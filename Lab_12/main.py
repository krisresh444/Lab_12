import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from openpyxl import Workbook
from bank_app.calculations import calculate_credit, calculate_installment, calculate_deposit

class ReportSaver:#   Класс для сохранения данных в текстовый (Word) или табличный (Excel) файл
    @staticmethod
    def save_to_docx(data: list, path: str, title: str = "Отчет"):
        doc = Document()
        doc.add_heading(title, level=1)
        for row in data:
            line = ", ".join([f"{key}: {value}" for key, value in row.items()])
            doc.add_paragraph(line)
        doc.save(path)

    @staticmethod
    def save_to_xlsx(data: list, path: str):
        wb = Workbook()
        ws = wb.active
        ws.append(list(data[0].keys()))  # Заголовок из ключей
        for row in data:
            ws.append(list(row.values()))  # Данные
        wb.save(path)

class BankingApp:#Класс для запуска графического интерфейса приложения
    def __init__(self):
        self.calculation_data = []
        self.root = tk.Tk()
        self.root.title("Банковские услуги")
        self._setup_ui()

    def _setup_ui(self):# Форма ввода данных
        ttk.Label(self.root, text="Выберите услугу:").grid(row=0, column=0, padx=5, pady=5)
        self.var_type = tk.StringVar(value="Кредит")
        ttk.Combobox(self.root, textvariable=self.var_type, values=["Кредит", "Рассрочка", "Вклад"]).grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(self.root, text="Сумма (руб.):").grid(row=1, column=0, padx=5, pady=5)
        self.entry_principal = ttk.Entry(self.root)
        self.entry_principal.grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(self.root, text="Ставка (% годовых):").grid(row=2, column=0, padx=5, pady=5)
        self.entry_rate = ttk.Entry(self.root)
        self.entry_rate.grid(row=2, column=1, padx=5, pady=5)

        ttk.Label(self.root, text="Срок (мес.):").grid(row=3, column=0, padx=5, pady=5)
        self.entry_term = ttk.Entry(self.root)
        self.entry_term.grid(row=3, column=1, padx=5, pady=5)

        ttk.Button(self.root, text="Рассчитать", command=self.calculate_action).grid(row=4, column=0, padx=5, pady=5)
        ttk.Button(self.root, text="Сохранить отчет", command=self.save_report_action).grid(row=4, column=1, padx=5, pady=5)

        self.output = tk.StringVar()
        ttk.Label(self.root, textvariable=self.output, wraplength=400).grid(row=5, column=0, columnspan=2, padx=5, pady=5)

    def calculate_action(self):#Выполняет расчет в зависимости от выбранной услуги
        try:
            principal = float(self.entry_principal.get())
            annual_rate = float(self.entry_rate.get())
            months = int(self.entry_term.get())
            service_type = self.var_type.get()

            if service_type == "Кредит":
                result, payment = calculate_credit(principal, annual_rate, months)
                self.output.set(f"Ежемесячный платёж: {payment} руб.")
            elif service_type == "Рассрочка":
                result = calculate_installment(principal, months)
                self.output.set("Рассрочка без процентов рассчитана.")
            elif service_type == "Вклад":
                final_amount = calculate_deposit(principal, annual_rate, months)
                self.output.set(f"Итоговая сумма вклада: {final_amount} руб.")
                result = [{"Месяц": months, "Сумма": final_amount}]
            
            self.calculation_data = result
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректные значения.")

    def save_report_action(self): #Сохраняет результаты расчёта в файл Word или Excel
        if not self.calculation_data:
            messagebox.showerror("Ошибка", "Нет данных для сохранения!")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word (.docx)", "*.docx"), ("Excel (.xlsx)", "*.xlsx")])
        if file_path:
            if file_path.endswith(".docx"):
                ReportSaver.save_to_docx(self.calculation_data, file_path, title="Банковский отчет")
            elif file_path.endswith(".xlsx"):
                ReportSaver.save_to_xlsx(self.calculation_data, file_path)
            messagebox.showinfo("Успех", "Отчет успешно сохранен!")

    def run(self):#Запуск основного цикла Tkinter
        self.root.mainloop()

if __name__ == "__main__":
    app = BankingApp()
    app.run()